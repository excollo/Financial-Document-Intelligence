"""
Docx Builder — generates a comprehensive Word document report 
from the extracted section results.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
from app.core.logging import get_logger

logger = get_logger(__name__)


class DocxBuilder:
    """
    Constructs a professional DRHP Analysis report in Word format.
    Uses MDN-style formatting and includes all section results.
    """

    def build(self, job_details: Dict[str, Any], section_results: List[Dict[str, Any]]) -> bytes:
        """
        Create the Word document in memory.
        
        Args:
            job_details: Metadata about the job (company name, date, etc)
            section_results: List of result documents from MongoDB
            
        Returns:
            Bytes of the generated .docx file
        """
        doc = Document()
        
        # 1. Title Page / Header
        title = doc.add_heading("DRHP / RHP Intelligence Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        company = job_details.get("document_name", "Analysis")
        p = doc.add_paragraph()
        run = p.add_run(f"Subject: {company}")
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report Generated: {job_details.get('created_at', 'N/A')}")
        doc.add_page_break()
        
        # 2. Table of Contents Placeholder
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph("This report contains a data-driven analysis of the provided filing based on the configured Standard Operating Procedure (SOP).")
        
        # 3. Section Results
        # Sort results by the order in SOP config if possible, 
        # or just alphabetically/chronologically
        section_results.sort(key=lambda x: x.get("section_id", ""))
        
        for res in section_results:
            section_label = res.get("section_id", "Unknown Section").replace("_", " ").title()
            doc.add_heading(section_label, 1)
            
            markdown = res.get("markdown", "")
            if markdown:
                # Basic markdown to docx conversion
                # We can refine this to handle tables/bolding properly
                self._add_markdown_content(doc, markdown)
            else:
                doc.add_paragraph("No content extracted for this section.")
            
            doc.add_paragraph("\n") # Spacer

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_markdown_content(self, doc, markdown: str):
        """Rudimentary markdown-to-docx converter."""
        lines = markdown.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("# "):
                doc.add_heading(line[2:], 2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 3)
            elif line.startswith("### "):
                doc.add_heading(line[4:], 4)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif "|" in line and "---" not in line:
                # Placeholder for table support (more complex)
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)
